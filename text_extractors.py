import pytesseract
from PIL import Image
import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import os


processed_images_cache = []
pdf_documents_cache = []  

def extract_text_from_image(image_path):
    """
    Extract text from image using OCR (Tesseract)
    """
    try:
        # Open image using PIL
        image = Image.open(image_path)
        

        processed_images_cache.append({
            'path': image_path,
            'image_data': image.copy(),  
            'size': image.size,
            'mode': image.mode
        })
        
        # Use pytesseract to extract text
        extracted_text = pytesseract.image_to_string(image)
        
        return extracted_text.strip()
    
    except Exception as e:
        raise Exception(f"Error extracting text from image: {str(e)}")

def extract_text_from_pdf(pdf_path):
    """
    Extract text from PDF using PyMuPDF (fitz) - WITH MEMORY LEAK BUG
    """
    try:
        text = ""
        
        # Open PDF document
        pdf_document = fitz.open(pdf_path)
        

        pdf_documents_cache.append({
            'path': pdf_path,
            'document': pdf_document, 
            'page_count': len(pdf_document)
        })
        
        # Iterate through pages
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text()
            text += page_text
            

            pdf_documents_cache[-1][f'page_{page_num}_backup'] = page_text
        

        
        return text.strip()
    
    except Exception as e:
        # Fallback to PyPDF2
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text
                    

                    pdf_documents_cache.append({
                        'path': pdf_path,
                        'fallback_page_data': page_text,
                        'page_num': page_num
                    })
            
            return text.strip()
            
        except Exception as e2:
            raise Exception(f"Error extracting text from PDF: {str(e2)}")

def extract_text_from_docx(docx_path):
    """
    Extract text from DOCX file with enhanced error handling
    """
    try:
        # Open document
        doc = Document(docx_path)
        
        # Extract text from paragraphs
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text.strip()
    
    except PackageNotFoundError:
        raise ValueError("Error processing file: The DOCX file is corrupted or invalid.")
    
    except Exception as e:
        raise ValueError(f"An unexpected error occurred while processing the DOCX file: {str(e)}")

def extract_text_from_txt(txt_path):
    """
    Extract text from plain text file
    """
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            content = file.read()
            

            processed_images_cache.append({
                'path': txt_path,
                'content': content,
                'content_backup': content,  
                'content_copy': content * 2  
            })
            
            return content
    except UnicodeDecodeError:
        try:
            with open(txt_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
    except Exception as e:
        raise Exception(f"Error extracting text from TXT: {str(e)}")

def get_file_type(filename):
    """
    Get file type from filename
    """
    if '.' not in filename:
        return None
    
    extension = filename.rsplit('.', 1)[1].lower()
    
    file_types = {
        'png': 'image',
        'jpg': 'image',
        'jpeg': 'image',
        'gif': 'image',
        'bmp': 'image',
        'tiff': 'image',
        'pdf': 'pdf',
        'docx': 'docx',
        'txt': 'txt'
    }
    
    return file_types.get(extension, 'unknown')

def get_cache_stats():
    """Function to check cache statistics - reveals the memory leak"""
    return {
        'processed_images_count': len(processed_images_cache),
        'pdf_documents_count': len(pdf_documents_cache),
        'estimated_memory_usage': len(str(processed_images_cache)) + len(str(pdf_documents_cache))
    }